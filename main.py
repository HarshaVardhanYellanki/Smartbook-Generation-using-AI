# importing required libraries:
import os
import openai
import time
import requests
import urllib.parse
from docx import Document
from docx.shared import Inches
from bs4 import BeautifulSoup
from PIL import Image
import shutil
import aspose.words as aw
from tkinter import *
from PIL import ImageTk, Image
from tkinter import messagebox as mb
import tkinter.font as font
import tkinter as tk
# download all the necessary libraries

api_key = "your api key here of google developers" # replace with your api
APIKEY = 'your open ai api key' # replace with your api
openai.api_key = APIKEY
source_folder = 'D:/Python IDE/Smartbook generation using AI/' #replace with Project folder path
destination_folder = 'C:/Users/Gowrilatha/Downloads' #replace with Download folder path

doc = Document()


# code to scrape links
def get_video_links(keyword, api_key):
    base_url = "https://www.googleapis.com/youtube/v3/search"
    params = {
        "part": "id",
        "q": urllib.parse.quote(keyword),
        "type": "video",
        "key": api_key,
        "maxResults": 5,  # Change this value to get more or fewer results
    }

    try:
        response = requests.get(base_url, params=params)
        response.raise_for_status()
    except requests.exceptions.RequestException as err:
        print("Error occurred:", err)
        return []

    data = response.json()
    video_links = []

    for item in data["items"]:
        video_id = item["id"]["videoId"]
        video_links.append(f"https://www.youtube.com/watch?v={video_id}")

    return video_links


# code to scrape images
def extract_images(topic):
    # Create a directory to store the images
    os.makedirs(topic, exist_ok=True)

    # URL to search for images
    url = f"https://www.google.com/search?q={topic}&source=lnms&tbm=isch"

    # Send a GET request to the URL
    response = requests.get(url)

    # Create a BeautifulSoup object to parse the HTML content
    soup = BeautifulSoup(response.content, 'html.parser')

    # Find all <img> tags
    img_tags = soup.find_all('img')

    # Counter to keep track of the image number
    img_num = 1

    # Download and save each image
    for img in img_tags:
        try:
            img_url = img['src']

            # Send a GET request to download the image
            img_response = requests.get(img_url)

            # Check if the response is an image
            if img_response.headers.get('content-type', '').startswith('image'):
                # Save the image as a PNG file
                with open(f"{topic}/image{img_num}.png", "wb") as f:
                    f.write(img_response.content)

                img_num += 1
        except:
            continue


# Algorithm
def get_answer(question):
    prompt = f"Question: {question}\nAnswer:"

    response = openai.Completion.create(
        engine='text-davinci-003',
        prompt=prompt,
        max_tokens=500,  # Adjust the desired length of the answer
        n=2,  # Specify the number of responses to generate
        stop=None,  # You can provide a stopping condition if needed
        temperature=0.9  # Adjust the creativity level (higher values = more random)
    )

    if 'choices' in response and len(response['choices']) > 0:
        return response['choices'][0]['text'].strip()
    else:
        return ''





def get_wikipedia_headings(concept):
    url = f"https://en.wikipedia.org/w/api.php?action=parse&format=json&page={concept}"

    response = requests.get(url)
    data = response.json()

    if 'error' in data:
        # print(f"Error: {data['error']['info']}")
        return

    sections = data['parse']['sections']
    for section in sections:
        heading = section['line']
        headings.append(heading)

    return headings





# default text
def on_entry_click(event):
    """Function to handle the event when the user clicks on the entry widget."""
    if my_text.get() == "Hello user, What is the concept to be learnt?":
        my_text.delete(0, tk.END)  # Clear the default text
        my_text.config(fg='black')


# Button function for input

headings = []
def print_user_input():
    concept = my_text.get()
    extract_images(concept)
    headings = get_wikipedia_headings(concept)
    if not headings:
        headings = ["Definitions", "Advantages", "Disadvantages"]
        # You can add as many as constraints in list, if dynamic constraints are not fetched from web.
        # Taking user input
        # Prompts and Required file generation
    count = 0
    print("Building the required format")
    for i in range(len(headings)):
        question = "What about %s of %s" % (headings[i], concept)
        count = count + 1
        answer = get_answer(question)
        book = answer
        with open('%s.txt' % (concept), 'a') as file:
            if (count == 1):
                file.write('\t\t%s\n\n' % (concept))
                file.write(' %s \n' % (headings[i]))
            else:
                file.write('\n\n %s \n' % (headings[i]))
            for i in book:
                file.write(i)

            if (count % 3 == 0):
                time.sleep(61)



    video_links = get_video_links(concept, api_key)

    if video_links:
        with open('%s.txt' % (concept), 'a') as file:
            file.write("\n\n Video links related to %s :\n" % (concept))
            for link in video_links:
                file.write(link)
                file.write('\n')
            file.write("\n\n Images related to %s :\n" % (concept))
    else:
        with open('%s.txt' % (concept), 'a') as file:
            file.write("\n No video links found for %s : \n" % (concept))

    with open('%s.txt' % (concept), 'r') as file:
        lines = file.read()
    doc.add_paragraph(lines)

    image_files = [f for f in os.listdir(concept) if os.path.isfile(os.path.join(concept, f))]

    # Insert images between paragraphs
    for image_file in image_files:
        image_path = os.path.join(concept, image_file)
        doc.add_picture(image_path, width=Inches(4.0))  # Adjust width as needed

    # Document version
    doc.save("%s.docx" % (concept))

    # Epub version
    doc1 = aw.Document("%s.docx" % (concept))
    doc1.save("%s.epub" % (concept))
    mb.showinfo(title="Desired Format", message="Please select desired format!")

op = 0
def button1():
    global op
    op = 1
    mb.showinfo(title="Download file!", message="Click download to start your download!")


def button2():
    global op
    op = 2
    mb.showinfo(title="Download file!", message="Click download to start your download!")


def button3():
    op = 3
    mb.showinfo(title="Download file!", message="Click download to start your download!")


def downloadbutton():
    files = os.listdir(source_folder)
    for file in files:
        source_path = os.path.join(source_folder, file)
        if (op == 1):
            if os.path.isfile(source_path) and file.lower().endswith('.txt'):
                destination_path = os.path.join(destination_folder, file)
                shutil.move(source_path, destination_path)

        if (op == 2):
            if os.path.isfile(source_path) and file.lower().endswith('.docx'):
                destination_path = os.path.join(destination_folder, file)
                shutil.move(source_path, destination_path)
        if (op == 3):
            if os.path.isfile(source_path) and file.lower().endswith('.epub'):
                destination_path = os.path.join(destination_folder, file)
                shutil.move(source_path, destination_path)
    mb.showinfo(title="File downloaded!", message="The file is successfully downloaded!")


# GUI
# main window creation
root = tk.Tk()
root.title("AI Smart book")
width = root.winfo_screenwidth()
height = root.winfo_screenheight()
root.geometry("%dx%d" % (width, height))

# font style
myFont = font.Font(family='Helvetica', size=30, weight="bold")
myFont1 = font.Font(family='Helvetica', size=20, weight="bold")

# Background image
img = Image.open('AIE.jpg')
bg = ImageTk.PhotoImage(img)
lbl = Label(root, image=bg, bg='pink')
lbl.config(bg="black", fg="white")
lbl.place(x=0, y=0)

# Text on main window
lbl1 = Label(root, text="          SmartText Generator using AI          ", font=myFont)
lbl1.config(bg="white")
lbl1.place(x=400, y=10)


# Entry Widget ::::
default_text = "Hello user, What is the concept to be learnt?"
my_text = tk.Entry(root, width=40, justify=CENTER, bg="white", font=('Times', 20, 'bold'))
my_text.insert(0, default_text)
my_text.config(fg='grey')
my_text.bind('<FocusIn>', on_entry_click)
my_text.place(x=200, y=100)

# extraction of dynamic constraints



# Submit button
b = Button(root, text="Submit", bg='yellow', font=myFont1, command=print_user_input)
b.place(x=400, y=140)

lbl1 = Label(root, text="Select Desired File Format :         ", font=myFont1)
lbl1.config(bg="white")
lbl1.place(x=200, y=250)

b = Button(root, text="Text file", bg='salmon', font=myFont1, command=button1)
b.place(x=200, y=320)

b = Button(root, text="Word Document", bg='salmon', font=myFont1, command=button2)
b.place(x=200, y=400)

b = Button(root, text="EBook", bg='salmon', font=myFont1, command=button3)
b.place(x=200, y=480)

b = Button(root, text="Download file", bg='yellow', font=myFont1, command=downloadbutton)
b.place(x=350, y=600)

root.mainloop()

